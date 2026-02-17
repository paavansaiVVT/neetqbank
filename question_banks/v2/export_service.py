"""
Export service for generating PDF, Excel (.xlsx), and Word (.docx) files
from question bank data.
"""
from __future__ import annotations

import io
import re
import csv
import json
import logging
from typing import Any

logger = logging.getLogger("qbank_v2_export")


# ─── helpers ────────────────────────────────────────────────────────

def _strip_latex(text: str) -> str:
    """Strip LaTeX delimiters for plain-text contexts (PDF, DOCX)."""
    text = re.sub(r"\$\$(.*?)\$\$", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\$(.*?)\$", r"\1", text)
    text = re.sub(r"\\\\", " ", text)
    text = re.sub(r"\\(vec|frac|sin|cos|tan|theta|alpha|beta|gamma|delta|epsilon|omega|pi|mu|lambda|sigma|sqrt|cdot|times|div|neq|leq|geq|approx|infty|sum|int|partial|nabla|rightarrow|leftarrow|Rightarrow|Leftarrow)\b", r"\1", text)
    text = re.sub(r"\\text\{(.*?)\}", r"\1", text)
    text = re.sub(r"\{([^{}]*)\}", r"\1", text)
    return text.strip()


def _option_letter(i: int) -> str:
    return chr(65 + i)


def _pdf_safe(text: str) -> str:
    """Replace Unicode chars unsupported by Helvetica with ASCII equivalents."""
    replacements = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u2022": "*",    # bullet
        "\u2713": "*",    # checkmark
        "\u2714": "*",    # heavy checkmark
        "\u2715": "x",    # multiplication x
        "\u00b0": "deg",  # degree symbol
        "\u00d7": "x",    # multiplication sign
        "\u2264": "<=",   # less than or equal
        "\u2265": ">=",   # greater than or equal
        "\u2260": "!=",   # not equal
        "\u221e": "inf",  # infinity
        "\u03b1": "alpha",
        "\u03b2": "beta",
        "\u03b3": "gamma",
        "\u03b4": "delta",
        "\u03b8": "theta",
        "\u03c0": "pi",
        "\u03c3": "sigma",
        "\u03bb": "lambda",
        "\u03bc": "mu",
        "\u03c9": "omega",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    # Fallback: encode to latin-1 and drop anything that can't be encoded
    return text.encode("latin-1", "replace").decode("latin-1")


# ─── JSON export ────────────────────────────────────────────────────

def export_json(
    items: list[dict[str, Any]],
    include_explanations: bool = True,
    include_metadata: bool = False,
    job_id: str = "library",
) -> tuple[io.BytesIO, str, str]:
    """Return (BytesIO, filename, media_type) for JSON export."""
    export_data = []
    for item in items:
        q: dict[str, Any] = {
            "question": item["question"],
            "options": item.get("options", []),
            "correct_answer": item["correct_answer"],
        }
        if include_explanations:
            q["explanation"] = item.get("explanation", "")
        if include_metadata:
            q["difficulty"] = item.get("difficulty", "medium")
            q["cognitive_level"] = item.get("cognitive_level", "")
            q["estimated_time"] = item.get("estimated_time", 0)
        export_data.append(q)

    payload = json.dumps({"questions": export_data, "count": len(export_data)}, indent=2, ensure_ascii=False)
    buf = io.BytesIO(payload.encode("utf-8"))
    return buf, f"qbank_{job_id}.json", "application/json"


# ─── CSV / Excel export ────────────────────────────────────────────

def export_excel(
    items: list[dict[str, Any]],
    include_explanations: bool = True,
    include_metadata: bool = False,
    job_id: str = "library",
) -> tuple[io.BytesIO, str, str]:
    """Return (BytesIO, filename, media_type) for a proper .xlsx file."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Questions"

    # ── Styles ──
    header_font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )
    cell_align = Alignment(vertical="top", wrap_text=True)

    # ── Header ──
    headers = ["#", "Question", "Option A", "Option B", "Option C", "Option D", "Correct Answer"]
    if include_explanations:
        headers.append("Explanation")
    if include_metadata:
        headers.extend(["Difficulty", "Cognitive Level", "Est. Time (min)"])

    for col, hdr in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=hdr)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # ── Data rows ──
    for row_idx, item in enumerate(items, 2):
        opts = item.get("options", [])
        q_text = _strip_latex(item.get("question", ""))
        values = [
            row_idx - 1,
            q_text,
            _strip_latex(opts[0]) if len(opts) > 0 else "",
            _strip_latex(opts[1]) if len(opts) > 1 else "",
            _strip_latex(opts[2]) if len(opts) > 2 else "",
            _strip_latex(opts[3]) if len(opts) > 3 else "",
            _strip_latex(item.get("correct_answer", "")),
        ]
        if include_explanations:
            values.append(_strip_latex(item.get("explanation", "")))
        if include_metadata:
            values.extend([
                item.get("difficulty", "medium"),
                item.get("cognitive_level", ""),
                str(item.get("estimated_time", "")),
            ])

        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col, value=val)
            cell.alignment = cell_align
            cell.border = thin_border

        # ── Difficulty color ──
        if include_metadata:
            diff_col = headers.index("Difficulty") + 1
            diff_cell = ws.cell(row=row_idx, column=diff_col)
            diff = item.get("difficulty", "medium")
            color_map = {
                "easy": "DCFCE7",
                "medium": "FEF3C7",
                "hard": "FEE2E2",
                "veryhard": "EDE9FE",
            }
            if diff in color_map:
                diff_cell.fill = PatternFill(start_color=color_map[diff], end_color=color_map[diff], fill_type="solid")

    # ── Column widths ──
    widths = {"#": 5, "Question": 50, "Correct Answer": 20, "Explanation": 40}
    for col_idx, hdr in enumerate(headers, 1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = widths.get(hdr, 18)

    # ── Freeze pane ──
    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf, f"qbank_{job_id}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# ─── PDF export ─────────────────────────────────────────────────────

def export_pdf(
    items: list[dict[str, Any]],
    include_explanations: bool = True,
    include_metadata: bool = False,
    job_id: str = "library",
    title: str = "Question Bank Export",
) -> tuple[io.BytesIO, str, str]:
    """Return (BytesIO, filename, media_type) for a proper PDF file."""
    from fpdf import FPDF

    class QBankPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 14)
            self.set_text_color(37, 99, 235)  # Primary blue
            self.cell(0, 10, _pdf_safe(title), ln=True, align="C")
            self.set_font("Helvetica", "", 9)
            self.set_text_color(107, 114, 128)
            self.cell(0, 6, f"{len(items)} questions", ln=True, align="C")
            self.ln(4)
            # Header line
            self.set_draw_color(229, 231, 235)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(156, 163, 175)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = QBankPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for i, item in enumerate(items):
        q_text = _pdf_safe(_strip_latex(item.get("question", "")))
        opts = item.get("options", [])
        correct = _pdf_safe(_strip_latex(item.get("correct_answer", "")))

        # Check if we need a new page (estimate height needed)
        needed = 40
        if include_explanations and item.get("explanation"):
            needed += 20
        if pdf.get_y() + needed > pdf.h - 25:
            pdf.add_page()

        # ── Question number + text ──
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(31, 41, 55)
        pdf.cell(12, 7, f"Q{i + 1}.")

        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(55, 65, 81)
        # Multi-cell for wrapping
        x_start = pdf.get_x()
        y_start = pdf.get_y()
        pdf.multi_cell(0, 6, q_text)
        pdf.ln(2)

        # ── Options ──
        for j, opt in enumerate(opts):
            opt_text = _pdf_safe(_strip_latex(opt))
            letter = _option_letter(j)
            is_correct = (opt_text == correct or letter == correct)

            if is_correct:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(22, 163, 74)  # Green
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(75, 85, 99)

            marker = " *" if is_correct else ""
            indent = 18
            pdf.set_x(pdf.l_margin + indent)
            remaining_w = pdf.w - pdf.l_margin - pdf.r_margin - indent
            pdf.multi_cell(remaining_w, 6, f"{letter}) {opt_text}{marker}")

        pdf.ln(1)

        # ── Correct answer line ──
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(22, 163, 74)
        pdf.set_x(pdf.l_margin + 18)
        pdf.cell(0, 5, f"Answer: {correct}", ln=True)

        # ── Explanation ──
        if include_explanations and item.get("explanation"):
            exp_text = _pdf_safe(_strip_latex(item["explanation"]))
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(107, 114, 128)
            indent = 18
            pdf.set_x(pdf.l_margin + indent)
            remaining_w = pdf.w - pdf.l_margin - pdf.r_margin - indent
            pdf.multi_cell(remaining_w, 5, f"Explanation: {exp_text}")

        # ── Metadata tags ──
        if include_metadata:
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(156, 163, 175)
            tags = []
            if item.get("difficulty"):
                tags.append(f"Difficulty: {item['difficulty']}")
            if item.get("cognitive_level"):
                tags.append(f"Cognitive: {item['cognitive_level']}")
            if item.get("estimated_time"):
                tags.append(f"Time: {item['estimated_time']} min")
            if tags:
                pdf.set_x(pdf.l_margin + 18)
                pdf.cell(0, 5, " | ".join(tags), ln=True)

        # ── Separator ──
        pdf.ln(3)
        pdf.set_draw_color(229, 231, 235)
        pdf.line(15, pdf.get_y(), pdf.w - 15, pdf.get_y())
        pdf.ln(4)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf, f"qbank_{job_id}.pdf", "application/pdf"


# ─── Word / DOCX export ────────────────────────────────────────────

def export_docx(
    items: list[dict[str, Any]],
    include_explanations: bool = True,
    include_metadata: bool = False,
    job_id: str = "library",
    title: str = "Question Bank Export",
) -> tuple[io.BytesIO, str, str]:
    """Return (BytesIO, filename, media_type) for a .docx file."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{len(items)} questions")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()  # spacer

    for i, item in enumerate(items):
        q_text = _strip_latex(item.get("question", ""))
        opts = item.get("options", [])
        correct = _strip_latex(item.get("correct_answer", ""))

        # ── Question ──
        q_para = doc.add_paragraph()
        q_num = q_para.add_run(f"Q{i + 1}. ")
        q_num.bold = True
        q_num.font.size = Pt(11)
        q_num.font.color.rgb = RGBColor(31, 41, 55)

        q_body = q_para.add_run(q_text)
        q_body.font.size = Pt(11)
        q_body.font.color.rgb = RGBColor(55, 65, 81)

        # ── Options ──
        for j, opt in enumerate(opts):
            opt_text = _strip_latex(opt)
            letter = _option_letter(j)
            is_correct = (opt_text == correct or letter == correct)

            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Inches(0.4)
            opt_para.paragraph_format.space_before = Pt(2)
            opt_para.paragraph_format.space_after = Pt(2)

            prefix = opt_para.add_run(f"{letter}) ")
            prefix.bold = is_correct
            prefix.font.size = Pt(10)

            text_run = opt_para.add_run(opt_text)
            text_run.font.size = Pt(10)

            if is_correct:
                prefix.font.color.rgb = RGBColor(22, 163, 74)
                text_run.font.color.rgb = RGBColor(22, 163, 74)
                text_run.bold = True
                check = opt_para.add_run("  ✓")
                check.font.color.rgb = RGBColor(22, 163, 74)
                check.bold = True
            else:
                prefix.font.color.rgb = RGBColor(75, 85, 99)
                text_run.font.color.rgb = RGBColor(75, 85, 99)

        # ── Answer line ──
        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.left_indent = Inches(0.4)
        ans_run = ans_para.add_run(f"Answer: {correct}")
        ans_run.bold = True
        ans_run.font.size = Pt(9)
        ans_run.font.color.rgb = RGBColor(22, 163, 74)

        # ── Explanation ──
        if include_explanations and item.get("explanation"):
            exp_text = _strip_latex(item["explanation"])
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.left_indent = Inches(0.4)
            exp_label = exp_para.add_run("Explanation: ")
            exp_label.bold = True
            exp_label.font.size = Pt(9)
            exp_label.font.color.rgb = RGBColor(107, 114, 128)
            exp_body = exp_para.add_run(exp_text)
            exp_body.font.size = Pt(9)
            exp_body.font.color.rgb = RGBColor(107, 114, 128)
            exp_body.italic = True

        # ── Metadata ──
        if include_metadata:
            tags = []
            if item.get("difficulty"):
                tags.append(f"Difficulty: {item['difficulty']}")
            if item.get("cognitive_level"):
                tags.append(f"Cognitive: {item['cognitive_level']}")
            if item.get("estimated_time"):
                tags.append(f"Time: {item['estimated_time']} min")
            if tags:
                meta_para = doc.add_paragraph()
                meta_para.paragraph_format.left_indent = Inches(0.4)
                meta_run = meta_para.add_run(" | ".join(tags))
                meta_run.font.size = Pt(8)
                meta_run.font.color.rgb = RGBColor(156, 163, 175)

        # ── Separator ──
        if i < len(items) - 1:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(4)
            border_run = sep.add_run("─" * 60)
            border_run.font.size = Pt(6)
            border_run.font.color.rgb = RGBColor(229, 231, 235)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, f"qbank_{job_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ─── Dispatcher ─────────────────────────────────────────────────────

def generate_export(
    items: list[dict[str, Any]],
    format: str,
    include_explanations: bool = True,
    include_metadata: bool = False,
    job_id: str = "library",
    title: str = "Question Bank Export",
) -> tuple[io.BytesIO, str, str]:
    """Dispatch to the appropriate exporter. Returns (buffer, filename, media_type)."""
    if format == "pdf":
        return export_pdf(items, include_explanations, include_metadata, job_id, title)
    elif format == "excel":
        return export_excel(items, include_explanations, include_metadata, job_id)
    elif format == "docx":
        return export_docx(items, include_explanations, include_metadata, job_id, title)
    elif format == "json":
        return export_json(items, include_explanations, include_metadata, job_id)
    else:
        raise ValueError(f"Unsupported export format: {format}")
